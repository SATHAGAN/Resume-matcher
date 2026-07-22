"""Step 1: Document Ingestion - pull raw text out of uploaded PDF/DOCX files."""
import fitz  # PyMuPDF
import docx


class UnsupportedFileType(ValueError):
    pass


def _extract_pdf_text(file_stream) -> str:
    file_stream.seek(0)
    doc = fitz.open(stream=file_stream.read(), filetype="pdf")
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _extract_docx_text(file_stream) -> str:
    file_stream.seek(0)
    document = docx.Document(file_stream)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(file_storage) -> str:
    """
    file_storage: a werkzeug FileStorage from request.files[...]
    Returns cleaned raw text, or raises UnsupportedFileType.
    """
    filename = (file_storage.filename or "").lower()

    if filename.endswith(".pdf"):
        text = _extract_pdf_text(file_storage.stream)
    elif filename.endswith(".docx"):
        text = _extract_docx_text(file_storage.stream)
    else:
        raise UnsupportedFileType(f"Unsupported file type: {file_storage.filename}")

    # Collapse excessive blank lines/whitespace from PDF extraction noise.
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    cleaned = "\n".join(lines)

    if not cleaned.strip():
        raise ValueError(f"No extractable text found in {file_storage.filename}")

    return cleaned
